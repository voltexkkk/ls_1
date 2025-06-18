import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.express as px
from datetime import datetime, timedelta
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib import rcParams

# ─── 페이지 설정 ────────────────────────────────────────
st.set_page_config(
    page_title="통합 전력 분석 시스템", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# matplotlib 한글 폰트 설정
plt.rcParams['font.family'] = ['Malgun Gothic', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# ─── 현대적인 CSS 스타일링 ────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;600;700&display=swap');
    
    * {
        font-family: 'Noto Sans KR', sans-serif;
    }
    
    .stApp {
        background-color: #fafbfc;
    }
    
    /* 메인 헤더 */
    .main-header {
        background: white;
        padding: 2.5rem;
        border-radius: 16px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
        margin-bottom: 2rem;
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 4px;
        background: linear-gradient(90deg, #1a73e8, #4285f4, #34a853);
    }
    
    .main-header h1 {
        color: #1a73e8;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: -1px;
    }
    
    .main-header .subtitle {
        color: #5f6368;
        font-size: 1.1rem;
        margin-top: 0.8rem;
        font-weight: 400;
    }
    
    /* 주요 지표 카드 */
    .main-metrics-card {
        background: white;
        border-radius: 16px;
        padding: 2rem;
        margin: 2rem 0;
        box-shadow: 0 2px 12px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
    }
    
    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1.5rem;
        margin-top: 1.5rem;
    }
    
    .metric-item {
        background: linear-gradient(145deg, #f8f9fa, #ffffff);
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
        border: 1px solid #e9ecef;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .metric-item::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 3px;
        background: linear-gradient(90deg, #1a73e8, #4285f4);
    }
    
    .metric-item:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(26, 115, 232, 0.15);
        background: linear-gradient(145deg, #ffffff, #f8f9fa);
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #5f6368;
        margin-bottom: 0.8rem;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .metric-value {
        font-size: 1.8rem;
        font-weight: 700;
        color: #202124;
        margin-bottom: 0.5rem;
        line-height: 1;
    }
    
    /* 비교 테이블 */
    .comparison-table {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
        margin: 1rem 0;
    }
    
    /* 섹션 헤더 */
    .section-header {
        background: white;
        color: #202124;
        padding: 1.5rem 2rem;
        border-radius: 12px;
        text-align: center;
        margin: 2rem 0 1rem 0;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border-left: 4px solid #1a73e8;
        position: relative;
    }
    
    .section-header h2 {
        margin: 0;
        font-size: 1.4rem;
        font-weight: 600;
    }
    
    /* 시간대별 카드 */
    .time-period-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        color: #202124;
        text-align: center;
        box-shadow: 0 3px 15px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
        transition: all 0.3s ease;
        border: 1px solid #e8eaed;
        position: relative;
        overflow: hidden;
    }
    
    .time-period-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
    }
    
    .daytime-card {
        border-left: 4px solid #f39c12;
    }
    
    .daytime-card::before {
        content: '';
        position: absolute;
        top: 0;
        right: 0;
        width: 60px;
        height: 60px;
        background: linear-gradient(135deg, #f39c12, #e67e22);
        border-radius: 0 12px 0 50px;
        opacity: 0.1;
    }
    
    .nighttime-card {
        border-left: 4px solid #34495e;
    }
    
    .nighttime-card::before {
        content: '';
        position: absolute;
        top: 0;
        right: 0;
        width: 60px;
        height: 60px;
        background: linear-gradient(135deg, #34495e, #2c3e50);
        border-radius: 0 12px 0 50px;
        opacity: 0.1;
    }
    
    .card-title {
        font-size: 1rem;
        font-weight: 600;
        margin-bottom: 0.8rem;
        color: #202124;
    }
    
    .card-value {
        font-size: 1.6rem;
        font-weight: 700;
        margin: 0.8rem 0;
        color: #202124;
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 0.5rem;
    }
    
    .card-change {
        font-size: 0.9rem;
        margin: 0.5rem 0;
        font-weight: 500;
        color: #5f6368;
    }
    
    .card-period {
        font-size: 0.75rem;
        color: #80868b;
        margin-top: 0.8rem;
        font-weight: 400;
    }
    
    .traffic-light {
        font-size: 1.2rem;
    }
    
    /* 차트 컨테이너 */
    .chart-container {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
    }
    
    /* 사이드바 스타일 */
    .sidebar-section {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
        margin-bottom: 1rem;
    }
    
    .sidebar-title {
        font-size: 1.1rem;
        color: #202124;
        font-weight: 600;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #f1f3f4;
    }
    
    /* 버튼 스타일 */
    .stButton > button {
        background: linear-gradient(135deg, #1a73e8 0%, #4285f4 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.7rem 1.5rem;
        font-weight: 600;
        transition: all 0.3s ease;
        width: 100%;
        margin-bottom: 0.5rem;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 15px rgba(26, 115, 232, 0.4);
        background: linear-gradient(135deg, #1557b0 0%, #3367d6 100%);
    }
    
    /* 선택박스 스타일 */
    .stSelectbox > div > div {
        background: white;
        border: 2px solid #e8eaed;
        border-radius: 8px;
        transition: all 0.2s ease;
    }
    
    .stSelectbox > div > div:focus-within {
        border-color: #1a73e8;
        box-shadow: 0 0 0 3px rgba(26, 115, 232, 0.1);
    }
    
    /* 날짜 입력 스타일 */
    .stDateInput > div > div > input {
        border: 2px solid #e8eaed;
        border-radius: 8px;
        padding: 0.5rem;
        transition: all 0.2s ease;
    }
    
    .stDateInput > div > div > input:focus {
        border-color: #1a73e8;
        box-shadow: 0 0 0 3px rgba(26, 115, 232, 0.1);
    }
    
    /* 데이터프레임 스타일 */
    .stDataFrame {
        background: white;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    /* Plotly 차트 스타일 */
    .js-plotly-plot {
        border-radius: 8px;
        overflow: hidden;
    }
    
    /* 경고 및 정보 메시지 */
    .stAlert {
        border-radius: 8px;
        border: none;
        box-shadow: 0 2px 10px rgba(0,0,0,0.08);
    }
    
    /* 푸터 스타일 */
    .footer-tips {
        background: white;
        border-radius: 16px;
        padding: 2rem;
        margin-top: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        border: 1px solid #e8eaed;
    }
    
    .tips-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 1.5rem;
        margin-top: 1.5rem;
    }
    
    .tip-card {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 12px;
        border-left: 4px solid #1a73e8;
        transition: all 0.2s ease;
    }
    
    .tip-card:hover {
        background: #f1f3f4;
        transform: translateY(-1px);
    }
    
    .tip-title {
        font-weight: 600;
        color: #202124;
        margin-bottom: 0.5rem;
        font-size: 1rem;
    }
    
    .tip-content {
        font-size: 0.85rem;
        color: #5f6368;
        line-height: 1.4;
    }
    
    /* 숨기고 싶은 요소 */
    .stDeployButton {
        display: none;
    }
    
    #MainMenu {
        visibility: hidden;
    }
    
    footer {
        visibility: hidden;
    }
    
    header {
        visibility: hidden;
    }
    
    /* 반응형 디자인 */
    @media (max-width: 768px) {
        .metrics-grid {
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
        }
        
        .main-header h1 {
            font-size: 2rem;
        }
        
        .metric-value {
            font-size: 1.5rem;
        }
        
        .tips-grid {
            grid-template-columns: 1fr;
            gap: 1rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# ========== 1. 데이터 로드 함수 ==========
@st.cache_data
def load_data():
    """데이터 로드 및 전처리"""
    try:
        df = pd.read_csv("./data/train.csv")
        df["측정일시"] = pd.to_datetime(df["측정일시"])
        df["날짜"] = df["측정일시"].dt.date
        df["시간"] = df["측정일시"].dt.hour
        df["월"] = df["측정일시"].dt.month
        df["일"] = df["측정일시"].dt.day
        df["년월"] = df["측정일시"].dt.to_period("M")
        return df
    except Exception as e:
        st.error(f"데이터 로드 중 오류 발생: {e}")
        return None

# ========== 2. 차트 생성 함수들 ==========
def create_matplotlib_chart(data, chart_type="line", title="Chart", xlabel="X", ylabel="Y", figsize=(10, 6)):
    """matplotlib로 간단한 차트 생성"""
    fig, ax = plt.subplots(figsize=figsize)
    fig.patch.set_facecolor('white')
    
    if chart_type == "line" and len(data.columns) >= 2:
        x_data = data.iloc[:, 0]
        y_data = data.iloc[:, 1]
        ax.plot(x_data, y_data, marker='o', linewidth=3, markersize=8, 
                color='#1a73e8', markerfacecolor='#4285f4', markeredgecolor='white', markeredgewidth=2)
    
    elif chart_type == "bar" and len(data.columns) >= 2:
        x_data = data.iloc[:, 0]
        y_data = data.iloc[:, 1]
        colors = ['#1a73e8', '#4285f4', '#34a853', '#ea4335']
        bars = ax.bar(x_data, y_data, color=colors[:len(x_data)], alpha=0.8, edgecolor='white', linewidth=2)
        
        # 막대 위에 값 표시
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                   f'{height:,.0f}', ha='center', va='bottom', fontweight='bold', color='#202124')
    
    elif chart_type == "pie" and len(data.columns) >= 2:
        labels = data.iloc[:, 0]
        sizes = data.iloc[:, 1]
        colors = ['#1a73e8', '#4285f4', '#34a853', '#fbbc04', '#ea4335', '#9aa0a6']
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', 
                                         colors=colors, startangle=90, 
                                         wedgeprops=dict(edgecolor='white', linewidth=2))
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
        ax.axis('equal')
    
    ax.set_title(title, fontsize=16, fontweight='bold', pad=20, color='#202124')
    ax.set_xlabel(xlabel, fontsize=12, fontweight='500', color='#202124')
    ax.set_ylabel(ylabel, fontsize=12, fontweight='500', color='#202124')
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e8eaed')
    ax.spines['bottom'].set_color('#e8eaed')
    
    plt.tight_layout()
    
    # 이미지로 변환
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer

def create_dual_axis_chart(df, x_col, y1_col, y2_col, title, x_title, y1_title, y2_title, add_time_zones=False):
    """듀얼 축 차트 생성"""
    fig = make_subplots(specs=[[{"secondary_y": True}]])

    if add_time_zones and x_title == "시간":
        fig.add_vrect(x0=-0.5, x1=5.5, fillcolor="rgba(26, 115, 232, 0.1)", layer="below", line_width=0,
                     annotation_text="야간 (00-05시)", annotation_position="top left")
        fig.add_vrect(x0=5.5, x1=17.5, fillcolor="rgba(251, 188, 4, 0.05)", layer="below", line_width=0,
                     annotation_text="주간 (06-17시)", annotation_position="top")
        fig.add_vrect(x0=17.5, x1=21.5, fillcolor="rgba(234, 67, 53, 0.08)", layer="below", line_width=0,
                     annotation_text="저녁 (18-21시)", annotation_position="top right")
        fig.add_vrect(x0=21.5, x1=23.5, fillcolor="rgba(26, 115, 232, 0.1)", layer="below", line_width=0,
                     annotation_text="야간 (22-23시)", annotation_position="top right")

    fig.add_trace(go.Scatter(x=df[x_col], y=df[y1_col], name=y1_title, 
                           line=dict(color="#1a73e8", width=4),
                           mode="lines+markers", 
                           marker=dict(size=8, color="#1a73e8", line=dict(color="white", width=2))), 
                 secondary_y=False)
    
    fig.add_trace(go.Scatter(x=df[x_col], y=df[y2_col], name=y2_title, 
                           line=dict(color="#ea4335", width=4),
                           mode="lines+markers", 
                           marker=dict(size=8, color="#ea4335", line=dict(color="white", width=2))), 
                 secondary_y=True)

    fig.update_xaxes(title_text=x_title, title_font=dict(size=14, color="#202124"))
    fig.update_yaxes(title_text=y1_title, secondary_y=False, title_font=dict(size=14, color="#1a73e8"))
    fig.update_yaxes(title_text=y2_title, secondary_y=True, title_font=dict(size=14, color="#ea4335"))
    
    fig.update_layout(
        title=dict(text=title, font=dict(size=18, color="#202124"), x=0.5),
        hovermode="x unified", 
        template="plotly_white", 
        height=500,
        font=dict(family="Noto Sans KR", size=12),
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="white",
        margin=dict(l=20, r=20, t=60, b=40)
    )
    return fig

def create_hourly_stack_chart(df):
    """시간별 스택 차트 생성"""
    hourly_worktype = df.groupby(["시간", "작업유형"])["전기요금(원)"].sum().unstack(fill_value=0)
    colors = {
        "Light_Load": "#34a853", 
        "Medium_Load": "#fbbc04", 
        "Maximum_Load": "#ea4335"
    }
    name_mapping = {
        "Light_Load": "경부하",
        "Medium_Load": "중간부하", 
        "Maximum_Load": "최대부하"
    }

    fig = go.Figure()
    for work_type in ["Light_Load", "Medium_Load", "Maximum_Load"]:
        if work_type in hourly_worktype.columns:
            fig.add_trace(go.Bar(
                name=name_mapping[work_type], 
                x=hourly_worktype.index, 
                y=hourly_worktype[work_type],
                marker_color=colors.get(work_type),
                marker_line=dict(color="white", width=1)
            ))

    fig.update_layout(
        barmode="stack", 
        title=dict(text="시간대별 작업유형별 전기요금 현황", font=dict(size=18, color="#202124"), x=0.5),
        xaxis_title="시간 (Hour)",
        yaxis_title="전기요금 (원)", 
        height=500, 
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="white",
        font=dict(family="Noto Sans KR", size=12),
        margin=dict(l=20, r=20, t=60, b=40)
    )
    return fig

def create_concentric_donut_chart(df):
    """도넛 차트 생성"""
    worktype_mwh = df.groupby("작업유형")["전력사용량(kWh)"].sum() / 1000
    total_mwh = worktype_mwh.sum()

    chart_data_map = {
        "Light_Load": {"name": "경부하", "color": "#34a853"},
        "Medium_Load": {"name": "중간부하", "color": "#fbbc04"},
        "Maximum_Load": {"name": "최대부하", "color": "#ea4335"}
    }

    labels, values, colors = [], [], []
    for work_type, data in chart_data_map.items():
        if work_type in worktype_mwh.index:
            labels.append(data["name"])
            values.append(worktype_mwh[work_type])
            colors.append(data["color"])

    fig = go.Figure(data=[go.Pie(
        labels=labels, 
        values=values, 
        hole=0.6, 
        marker=dict(colors=colors, line=dict(color="white", width=3)),
        pull=[0.05]*len(labels),
        textfont=dict(size=14, color="white"),
        textinfo="label+percent"
    )])
    
    fig.update_layout(
        title=dict(text="부하대별 에너지 사용량 분포", font=dict(size=18, color="#202124"), x=0.5),
        height=500,
        annotations=[dict(
            text=f"<b>총 {total_mwh:,.1f}</b><br>MWh", 
            x=0.5, y=0.5, 
            showarrow=False,
            font=dict(size=16, color="#202124")
        )],
        font=dict(family="Noto Sans KR"),
        paper_bgcolor="white",
        margin=dict(l=20, r=20, t=60, b=40)
    )
    return fig

# ========== 3. 카드 및 테이블 생성 함수들 ==========
def create_main_metrics_card(summary_data, period_label):
    """주요 지표 카드 생성"""
    if summary_data.empty:
        return ""
    
    total_kwh = summary_data["전력사용량(kWh)"].sum()
    total_cost = summary_data["전기요금(원)"].sum()
    total_carbon = summary_data["탄소배출량(tCO2)"].sum()
    avg_price = total_cost / total_kwh if total_kwh > 0 else 0
    
    card_html = f"""
    <div class="main-metrics-card">
        <h3 style="text-align: center; margin-bottom: 0; color: #202124; font-size: 1.4rem;">{period_label} 주요 지표</h3>
        <div class="metrics-grid">
            <div class="metric-item">
                <div class="metric-label">전력사용량</div>
                <div class="metric-value">{total_kwh:,.1f}</div>
                <div style="font-size: 0.8rem; color: #80868b; font-weight: 500;">kWh</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">전기요금</div>
                <div class="metric-value">₩{total_cost:,.0f}</div>
                <div style="font-size: 0.8rem; color: #80868b; font-weight: 500;">원</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">평균 단가</div>
                <div class="metric-value">{avg_price:.1f}</div>
                <div style="font-size: 0.8rem; color: #80868b; font-weight: 500;">원/kWh</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">탄소배출량</div>
                <div class="metric-value">{total_carbon:.2f}</div>
                <div style="font-size: 0.8rem; color: #80868b; font-weight: 500;">tCO₂</div>
            </div>
        </div>
    </div>
    """
    return card_html

def calculate_kepco_rate_impact(pf_value, time_period):
    """한전 요금 영향 계산"""
    if time_period == "daytime":
        adjusted_pf = max(60, min(95, pf_value))
        if adjusted_pf >= 90:
            rate_impact = -(adjusted_pf - 90) * 0.5
        else:
            rate_impact = (90 - adjusted_pf) * 0.5
    else:
        if pf_value <= 0:
            adjusted_pf = 100
        else:
            adjusted_pf = max(60, pf_value)
        if adjusted_pf >= 95:
            rate_impact = 0
        else:
            rate_impact = (95 - adjusted_pf) * 0.5
    return rate_impact

def get_traffic_light_and_message(current_pf, previous_pf, time_period):
    """신호등 및 메시지 생성"""
    current_impact = calculate_kepco_rate_impact(current_pf, time_period)
    previous_impact = calculate_kepco_rate_impact(previous_pf, time_period)
    rate_difference = current_impact - previous_impact
    
    if abs(rate_difference) < 0.1:
        traffic_light = "🟡"
        message = "전일과 동일"
    elif rate_difference > 0:
        traffic_light = "🔴"
        message = f"전일대비 +{rate_difference:.1f}% 더 냄"
    else:
        traffic_light = "🟢"
        message = f"전일대비 {rate_difference:.1f}% 덜 냄"
    return traffic_light, message

def create_simple_power_factor_card(period_name, icon, current_pf, previous_pf, time_period, card_class):
    """역률 카드 생성"""
    traffic_light, message = get_traffic_light_and_message(current_pf, previous_pf, time_period)
    pf_type = "지상" if time_period == "daytime" else "진상"
    time_range = "(09-23시)" if time_period == "daytime" else "(23-09시)"
    
    card_html = f"""
    <div class="time-period-card {card_class}">
        <div class="card-title">{icon} {period_name} 역률 {time_range}</div>
        <div class="card-value"><span class="traffic-light">{traffic_light}</span>{pf_type} {current_pf:.1f}%</div>
        <div class="card-change">{message}</div>
        <div class="card-period">한전 요금체계 기준</div>
    </div>
    """
    return card_html

def create_summary_table(current_data, period_type="일"):
    """요약 테이블 생성"""
    numeric_columns = [("전력사용량(kWh)", "kWh"), ("지상무효전력량(kVarh)", "kVarh"), ("진상무효전력량(kVarh)", "kVarh"),
                      ("탄소배출량(tCO2)", "tCO2"), ("지상역률(%)", "%"), ("진상역률(%)", "%"), ("전기요금(원)", "원")]
    ratio_cols = {"지상역률(%)", "진상역률(%)"}

    rows = []
    for col, unit in numeric_columns:
        if col in ratio_cols:
            val = current_data[col].mean()
        else:
            val = current_data[col].sum()
        name = col.split("(")[0]
        rows.append({"항목": name, f"현재{period_type} 값": f"{val:.2f}", "단위": unit})
    return pd.DataFrame(rows)

def create_comparison_table(current_data, previous_data, period_type="일"):
    """비교 테이블 생성"""
    comparison_dict = {"항목": [], f"현재{period_type}": [], f"이전{period_type}": [], "변화량": [], "변화율(%)": []}
    numeric_columns = ["전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"]

    for col in numeric_columns:
        if col in ["지상역률(%)", "진상역률(%)"]:
            current_val = current_data[col].mean()
            previous_val = previous_data[col].mean()
        else:
            current_val = current_data[col].sum()
            previous_val = previous_data[col].sum()

        change = current_val - previous_val
        change_pct = (change / previous_val * 100) if previous_val != 0 else 0

        comparison_dict["항목"].append(col)
        comparison_dict[f"현재{period_type}"].append(f"{current_val:.2f}")
        comparison_dict[f"이전{period_type}"].append(f"{previous_val:.2f}")
        comparison_dict["변화량"].append(f"{change:+.2f}")
        comparison_dict["변화율(%)"].append(f"{change_pct:+.1f}%")
    return pd.DataFrame(comparison_dict)

# ========== 4. 개선된 보고서 생성 함수 ==========
def create_comprehensive_docx_report_with_charts(df, current_data, daily_data, selected_date, view_type="월별", selected_month=1, period_label="전체"):
    """현재 화면 설정에 따른 동적 보고서 생성"""
    doc = Document()
    
    # 전체 문서에 테두리 추가
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = sectPr.xpath('.//w:pgBorders')
        if not pgBorders:
            from docx.oxml import parse_xml
            pgBorders_xml = '''
            <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                         w:offsetFrom="page">
                <w:top w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:left w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:bottom w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:right w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            </w:pgBorders>
            '''
            pgBorders = parse_xml(pgBorders_xml)
            sectPr.append(pgBorders)
    
    # 보고서 헤더 테이블
    header_table = doc.add_table(rows=6, cols=4)
    header_table.style = 'Table Grid'
    
    # 제목 행
    title_cell = header_table.rows[0].cells[0]
    title_cell.merge(header_table.rows[0].cells[3])
    title_cell.text = "보 고 서"
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para = title_cell.paragraphs[0]
    title_para.runs[0].font.size = Cm(0.8)
    title_para.runs[0].bold = True
    
    # 헤더 정보 입력
    header_data = [
        ("보고처", "에너지관리팀", "보고서명", f"전력 분석 보고서 ({period_label})"),
        ("장소", "본사", "취급분류", "○기밀 ●보통"),
        ("작성일자", datetime.now().strftime("%Y년 %m월 %d일"), "작성자", "전력분석팀"),
        ("참가자", "에너지관리팀, 시설관리팀, 경영진", "", ""),
        ("자료출처", "전력량계 실시간 데이터, 한국전력공사 요금체계", "", "")
    ]
    
    for i, (col1, val1, col2, val2) in enumerate(header_data, 1):
        header_table.rows[i].cells[0].text = col1
        header_table.rows[i].cells[1].text = val1
        if col2:
            header_table.rows[i].cells[2].text = col2
            header_table.rows[i].cells[3].text = val2
        else:
            header_table.rows[i].cells[1].merge(header_table.rows[i].cells[3])
    
    doc.add_paragraph()
    
    # 보고내용 제목
    content_title = doc.add_heading('보고내용', level=1)
    content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === 1. 기간별 분석 ===
    if view_type == "월별":
        doc.add_heading(f'1. {selected_month}월 전력 사용 분석', level=2)
        
        if not current_data.empty:
            total_kwh = current_data["전력사용량(kWh)"].sum()
            total_cost = current_data["전기요금(원)"].sum()
            avg_pf = current_data["지상역률(%)"].mean()
            total_carbon = current_data["탄소배출량(tCO2)"].sum()
            avg_price = total_cost / total_kwh if total_kwh > 0 else 0
            
            doc.add_paragraph(f"□ {selected_month}월 총 전력사용량: {total_kwh:,.1f} kWh")
            doc.add_paragraph(f"□ {selected_month}월 총 전기요금: {total_cost:,.0f} 원")
            doc.add_paragraph(f"□ {selected_month}월 평균 단가: {avg_price:.1f} 원/kWh")
            doc.add_paragraph(f"□ {selected_month}월 평균 역률: {avg_pf:.1f}%")
            doc.add_paragraph(f"□ {selected_month}월 탄소배출량: {total_carbon:.2f} tCO2")
    else:
        doc.add_heading(f'1. {period_label} 전력 사용 분석', level=2)
        
        if not current_data.empty:
            total_kwh = current_data["전력사용량(kWh)"].sum()
            total_cost = current_data["전기요금(원)"].sum()
            avg_pf = current_data["지상역률(%)"].mean()
            total_carbon = current_data["탄소배출량(tCO2)"].sum()
            avg_price = total_cost / total_kwh if total_kwh > 0 else 0
            
            doc.add_paragraph(f"□ 기간 총 전력사용량: {total_kwh:,.1f} kWh")
            doc.add_paragraph(f"□ 기간 총 전기요금: {total_cost:,.0f} 원")
            doc.add_paragraph(f"□ 기간 평균 단가: {avg_price:.1f} 원/kWh")
            doc.add_paragraph(f"□ 기간 평균 역률: {avg_pf:.1f}%")
            doc.add_paragraph(f"□ 기간 탄소배출량: {total_carbon:.2f} tCO2")
    
    return doc

# ========== 5. 메인 함수 ==========
def main():
    # 메인 헤더
    st.markdown("""
    <div class="main-header">
        <h1>통합 전력 분석 시스템</h1>
        <div class="subtitle">스마트 에너지 관리 및 효율 최적화 대시보드</div>
    </div>
    """, unsafe_allow_html=True)

    df = load_data()
    if df is None:
        st.stop()

    # ─── 사이드바 설정 ────────────────────────────────────────
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-title">분석 설정</div>
        </div>
        """, unsafe_allow_html=True)
        
        filtered_df = df.copy()
        date_range = (df["날짜"].min(), df["날짜"].max())
        work_types = df["작업유형"].unique()
        
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-title">상세 분석 옵션</div>
        </div>
        """, unsafe_allow_html=True)
        
        numeric_columns = ["전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"]
        col1_select = st.selectbox("첫 번째 분석 컬럼", numeric_columns, index=0)
        col2_select = st.selectbox("두 번째 분석 컬럼", numeric_columns, index=6)

        # 보고서 생성 옵션
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-title">보고서 생성</div>
        </div>
        """, unsafe_allow_html=True)
            
        if st.button("종합 보고서 생성", key="generate_complete_report"):
            with st.spinner("보고서 생성 중..."):
                try:
                    # 현재 설정된 분석 조건 가져오기
                    view_type = st.session_state.get('analysis_period', '월별')
                    
                    if view_type == "월별":
                        selected_month = st.session_state.get('month_selector', 1)
                        current_year = filtered_df["년월"].dt.year.max()
                        current_data = filtered_df[
                            (filtered_df["년월"].dt.year == current_year) &
                            (filtered_df["년월"].dt.month == selected_month)
                        ]
                        period_label = f"{selected_month}월"
                    else:
                        selected_range = st.session_state.get('period_range_selector', None)
                        if selected_range and len(selected_range) == 2:
                            start_day, end_day = selected_range
                            current_data = filtered_df[
                                (filtered_df["날짜"] >= start_day) & 
                                (filtered_df["날짜"] <= end_day)
                            ]
                            period_label = f"{start_day} ~ {end_day} 기간"
                        else:
                            current_data = filtered_df
                            period_label = "전체 기간"
                    
                    # 최근 날짜 데이터
                    latest_date = filtered_df["날짜"].max()
                    daily_data = filtered_df[filtered_df["날짜"] == latest_date]
                    
                    # 보고서 생성
                    doc = create_comprehensive_docx_report_with_charts(
                        filtered_df, current_data, daily_data, latest_date, 
                        view_type, selected_month if view_type == "월별" else None, period_label
                    )
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"전력분석보고서_{timestamp}.docx"
                    
                    st.download_button(
                        label="보고서 다운로드",
                        data=doc_buffer.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_complete_report"
                    )
                    
                    st.success("보고서 생성 완료!")
                    st.info(f"파일명: {filename}")
                    
                except Exception as e:
                    st.error(f"보고서 생성 중 오류 발생: {str(e)}")
                    st.info("오류가 지속되면 다른 날짜나 기간을 선택해보세요.")

    summary_data = filtered_df.copy()
    period_label = "전체"

    # ─── 필터링 옵션 ────────────────────────────────────────
    filter_col1, filter_col2, filter_col3, filter_col4 = st.columns([1, 0.8, 1.2, 1])
    
    with filter_col1:
        view_type = st.selectbox("분석 기간", ["월별", "일별"], key="analysis_period")
    
    with filter_col2:
        if view_type == "월별":
            current_year = filtered_df["년월"].dt.year.max()
            months = list(range(1, 13))
            default_month = filtered_df["년월"].dt.month.max()
            selected_month = st.selectbox("월", months, index=int(default_month) - 1, key="month_selector")
        else:
            st.markdown("")
    
    with filter_col3:
        if view_type == "일별":
            from datetime import date
            selected_range = st.date_input(
                "기간 선택", 
                value=(date(2024, 1, 1), date(2024, 1, 5)),
                min_value=date_range[0] if len(date_range) == 2 else df["날짜"].min(),
                max_value=date_range[1] if len(date_range) == 2 else df["날짜"].max(),
                key="period_range_selector"
            )
        else:
            st.markdown("")
    
    with filter_col4:
        st.markdown("")
    
    current_data = pd.DataFrame()
    previous_data = pd.DataFrame()

    # 데이터 처리 로직
    if view_type == "월별":
        current_data = filtered_df[(filtered_df["년월"].dt.year == current_year) & (filtered_df["년월"].dt.month == selected_month)]
        summary_data = current_data
        period_label = f"{selected_month}월"

        if selected_month > 1:
            prev_year = current_year
            prev_month = selected_month - 1
        else:
            prev_year = current_year - 1
            prev_month = 12
        previous_data = filtered_df[(filtered_df["년월"].dt.year == prev_year) & (filtered_df["년월"].dt.month == prev_month)]

    else:
        if not isinstance(selected_range, tuple) or len(selected_range) != 2:
            st.warning("📅 날짜 범위를 선택해주세요")
        else:
            start_day, end_day = selected_range
            if start_day > end_day:
                st.warning("⛔ 시작 날짜가 종료 날짜보다 이후입니다.")
            else:
                period_df = filtered_df[(filtered_df["날짜"] >= start_day) & (filtered_df["날짜"] <= end_day)]
                if period_df.empty:
                    st.info(f"{start_day} ~ {end_day} 구간에는 데이터가 없습니다.")
                else:
                    current_data = period_df
                    summary_data = period_df
                    period_label = f"{start_day} ~ {end_day} 기간"
                    
                    days = (end_day - start_day).days + 1
                    prev_start = start_day - timedelta(days=days)
                    prev_end = start_day - timedelta(days=1)
                    previous_data = filtered_df[(filtered_df["날짜"] >= prev_start) & (filtered_df["날짜"] <= prev_end)]

    # 주요 지표 카드
    if not summary_data.empty:
        main_metrics_card = create_main_metrics_card(summary_data, period_label)
        st.markdown(main_metrics_card, unsafe_allow_html=True)

    # ─── 차트 섹션 ────────────────────────────────────────
    chart_col1, chart_col2 = st.columns([2, 1])
    
    with chart_col1:
        st.markdown('<div class="chart-container">', unsafe_allow_html=True)
        if view_type == "월별":
            monthly_data = (filtered_df.groupby("년월").agg({
                col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
            }).reset_index())
            monthly_data["년월_str"] = monthly_data["년월"].astype(str)

            fig = create_dual_axis_chart(monthly_data, "년월_str", col1_select, col2_select,
                                       f"월별 {col1_select} vs {col2_select} 비교", "월", col1_select, col2_select)
            st.plotly_chart(fig, use_container_width=True)

        else:
            if isinstance(current_data, pd.DataFrame) and not current_data.empty:
                daily_data = (period_df.groupby("날짜").agg({
                    col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                    col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
                }).reset_index())

                fig = create_dual_axis_chart(daily_data, "날짜", col1_select, col2_select,
                                           f"{start_day} ~ {end_day} 날짜별 {col1_select} vs {col2_select}",
                                           "날짜", col1_select, col2_select)
                st.plotly_chart(fig, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with chart_col2:
        # 월별 분석일 때 비교 테이블
        if view_type == "월별" and not current_data.empty and not previous_data.empty:
            st.markdown('<div class="section-header" style="margin: 0 0 1rem 0;"><h3>전월 대비 분석</h3></div>', unsafe_allow_html=True)
            comparison_df = create_comparison_table(current_data, previous_data, "월")
            st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
            st.dataframe(comparison_df, use_container_width=True, hide_index=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.markdown("")

    st.markdown("---")

    # ─── 특정일 시간별 분석 ────────────────────────────────────────
    st.markdown('<div class="section-header"><h2>특정일 시간별 에너지 사용 분석</h2></div>', unsafe_allow_html=True)

    col1, col2 = st.columns([3, 1])
    daily_df = pd.DataFrame()

    with col1:
        available_dates = sorted(filtered_df["날짜"].unique())
        if available_dates:
            min_d, max_d = available_dates[0], available_dates[-1]
            default_d = max_d
            selected_date = st.date_input("분석할 날짜 선택", value=default_d, min_value=min_d, max_value=max_d, key="daily_date_selector")

            daily_df = filtered_df[filtered_df["날짜"] == selected_date]
            if daily_df.empty:
                st.warning(f"{selected_date} 데이터가 없습니다.")
            else:
                hourly_data = (daily_df.groupby("시간").agg({
                    col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                    col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
                }).reset_index())

                full_hours = pd.DataFrame({"시간": list(range(24))})
                hourly_data = pd.merge(full_hours, hourly_data, on="시간", how="left").fillna(0)

                st.markdown('<div class="chart-container">', unsafe_allow_html=True)
                fig = create_dual_axis_chart(hourly_data, "시간", col1_select, col2_select,
                                           f"{selected_date} 시간별 {col1_select} vs {col2_select} 비교",
                                           "시간", col1_select, col2_select, add_time_zones=True)

                fig.update_xaxes(tickmode="linear", tick0=0, dtick=1, title_text="시간")
                st.plotly_chart(fig, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("선택된 조건에 맞는 데이터가 없습니다.")

    with col2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown('<div class="section-header" style="margin: 0;"><h3>전일 대비 역률 요금</h3></div>', unsafe_allow_html=True)

        if available_dates and selected_date in available_dates:
            try:
                date_idx = available_dates.index(selected_date)
                if date_idx > 0:
                    previous_date = available_dates[date_idx - 1]
                    previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]

                    if not daily_df.empty and not previous_daily_df.empty:
                        current_daytime = daily_df[(daily_df['시간'] >= 9) & (daily_df['시간'] < 23)]
                        previous_daytime = previous_daily_df[(previous_daily_df['시간'] >= 9) & (previous_daily_df['시간'] < 23)]
                        current_nighttime = daily_df[(daily_df['시간'] >= 23) | (daily_df['시간'] < 9)]
                        previous_nighttime = previous_daily_df[(previous_daily_df['시간'] >= 23) | (previous_daily_df['시간'] < 9)]
                        
                        if len(current_daytime) > 0:
                            current_daytime_raw = current_daytime['지상역률(%)'].mean()
                            current_daytime_pf = max(60, min(95, current_daytime_raw))
                        else:
                            current_daytime_pf = 90
                        
                        if len(previous_daytime) > 0:
                            previous_daytime_raw = previous_daytime['지상역률(%)'].mean()
                            previous_daytime_pf = max(60, min(95, previous_daytime_raw))
                        else:
                            previous_daytime_pf = 90
                        
                        if len(current_nighttime) > 0:
                            current_leading_raw = current_nighttime['진상역률(%)'].mean()
                            if current_leading_raw > 0:
                                current_nighttime_pf = max(60, current_leading_raw)
                            else:
                                current_nighttime_pf = 100
                        else:
                            current_nighttime_pf = 100
                        
                        if len(previous_nighttime) > 0:
                            previous_leading_raw = previous_nighttime['진상역률(%)'].mean()
                            if previous_leading_raw > 0:
                                previous_nighttime_pf = max(60, previous_leading_raw)
                            else:
                                previous_nighttime_pf = 100
                        else:
                            previous_nighttime_pf = 100
                        
                        daytime_card = create_simple_power_factor_card("주간", "주간", current_daytime_pf, previous_daytime_pf, "daytime", "daytime-card")
                        nighttime_card = create_simple_power_factor_card("야간", "야간", current_nighttime_pf, previous_nighttime_pf, "nighttime", "nighttime-card")
                        
                        st.markdown(daytime_card, unsafe_allow_html=True)
                        st.markdown(nighttime_card, unsafe_allow_html=True)
                    else:
                        st.info("선택된 날짜 또는 전일 데이터가 없습니다.")
                else:
                    if not daily_df.empty:
                        summary_df = create_summary_table(daily_df, "일")
                        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                        st.dataframe(summary_df, use_container_width=True, hide_index=True)
                        st.markdown('</div>', unsafe_allow_html=True)
                        st.info("첫 번째 날짜로 전일 데이터가 없어 비교할 수 없습니다.")
                    else:
                        st.info("선택된 날짜의 데이터가 없습니다.")
            except (ValueError, IndexError):
                st.info("이전 날짜를 찾을 수 없습니다.")

    # 상세 비교 데이터 표
    if available_dates and selected_date in available_dates:
        try:
            date_idx = available_dates.index(selected_date)
            if date_idx > 0:
                previous_date = available_dates[date_idx - 1]
                previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]
                
                if not daily_df.empty and not previous_daily_df.empty:
                    st.markdown('<div class="section-header"><h3>상세 비교 데이터</h3></div>', unsafe_allow_html=True)
                    comparison_df = create_comparison_table(daily_df, previous_daily_df, "일")
                    st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                    st.dataframe(comparison_df, use_container_width=True, hide_index=True)
                    st.markdown('</div>', unsafe_allow_html=True)
        except (ValueError, IndexError):
            pass

    st.markdown("---")

    # ─── 시간대별 현황 차트 ────────────────────────────────────────
    if not daily_df.empty:
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            st.markdown('<div class="chart-container">', unsafe_allow_html=True)
            st.plotly_chart(create_hourly_stack_chart(daily_df), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
        with col_chart2:
            st.markdown('<div class="chart-container">', unsafe_allow_html=True)
            st.plotly_chart(create_concentric_donut_chart(daily_df), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown(f'<div class="section-header"><h3>{selected_date} 작업유형별 상세 분석</h3></div>', unsafe_allow_html=True)
        worktype_stats = (daily_df.groupby("작업유형").agg(
            전력사용량_합계=("전력사용량(kWh)", "sum"),
            전기요금_합계=("전기요금(원)", "sum"),
            평균_지상역률=("지상역률(%)", "mean"),
            탄소배출량_합계=("탄소배출량(tCO2)", "sum")
        ).round(2))
        
        # 작업유형 한글 변환
        worktype_stats.index = worktype_stats.index.map({
            'Light_Load': '경부하',
            'Medium_Load': '중간부하',
            'Maximum_Load': '최대부하'
        })
        
        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
        st.dataframe(worktype_stats, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            st.markdown('<div class="chart-container">', unsafe_allow_html=True)
            st.plotly_chart(create_hourly_stack_chart(filtered_df), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
        with col_chart2:
            st.markdown('<div class="chart-container">', unsafe_allow_html=True)
            st.plotly_chart(create_concentric_donut_chart(filtered_df), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-header"><h3>전체 기간 작업유형별 상세 분석</h3></div>', unsafe_allow_html=True)
        worktype_stats = (filtered_df.groupby("작업유형").agg(
            전력사용량_합계=("전력사용량(kWh)", "sum"),
            전기요금_합계=("전기요금(원)", "sum"),
            평균_지상역률=("지상역률(%)", "mean"),
            탄소배출량_합계=("탄소배출량(tCO2)", "sum")
        ).round(2))
        
        # 작업유형 한글 변환
        worktype_stats.index = worktype_stats.index.map({
            'Light_Load': '경부하',
            'Medium_Load': '중간부하',
            'Maximum_Load': '최대부하'
        })
        
        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
        st.dataframe(worktype_stats, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    # ─── 푸터 ────────────────────────────────────────
    st.markdown("---")
    st.markdown("""
    <div class="footer-tips">
        <h4 style="color: #1a73e8; margin-bottom: 1rem; text-align: center;">에너지 효율 개선 제안</h4>
        <div class="tips-grid">
            <div class="tip-card">
                <div class="tip-title">역률 개선</div>
                <div class="tip-content">진상/지상 역률을 95% 이상 유지하여 요금 할증을 방지하세요.</div>
            </div>
            <div class="tip-card">
                <div class="tip-title">부하 분산</div>
                <div class="tip-content">최대부하 시간대 사용량을 경부하 시간대로 이전하세요.</div>
            </div>
            <div class="tip-card">
                <div class="tip-title">탄소 절감</div>
                <div class="tip-content">에너지 효율 개선으로 탄소배출량을 줄이고 ESG 경영을 실천하세요.</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()